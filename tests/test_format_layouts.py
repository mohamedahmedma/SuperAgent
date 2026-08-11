"""A4 config (env-driven sizes/strategy) + DOCX/XLSX/HTML layout parsers.

Format tests author real files with python-docx / openpyxl / plain HTML, then run
them through the block parsers and the full DocumentLoader pipeline — true
end-to-end coverage without binary fixtures in the repo.
"""
import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import Mock, patch

import backend.indexing.document_loader as document_loader_module
from backend.indexing.document_loader import DocumentLoader, SentenceSplitter
from backend.indexing.docx_layout import parse_docx_blocks
from backend.indexing.html_layout import parse_html_blocks
from backend.indexing.xlsx_layout import parse_xlsx_blocks


class ChunkConfigTests(unittest.TestCase):
    def test_chunk_size_env_drives_level_sizes(self):
        with patch.dict(os.environ, {"CHUNK_SIZE": "1000", "CHUNK_OVERLAP": "120"}):
            loader = DocumentLoader()
        self.assertEqual(3000, loader._level_1_size)
        self.assertEqual(2000, loader._level_2_size)
        self.assertEqual(1000, loader._level_3_size)

    def test_per_level_size_overrides(self):
        with patch.dict(os.environ, {"CHUNK_L1_SIZE": "5000", "CHUNK_L2_SIZE": "2500"}):
            loader = DocumentLoader()
        self.assertEqual(5000, loader._level_1_size)
        self.assertEqual(2500, loader._level_2_size)

    def test_invalid_strategy_falls_back_to_recursive(self):
        with patch.dict(os.environ, {"CHUNK_STRATEGY": "quantum"}):
            loader = DocumentLoader()
        self.assertEqual("recursive", loader._strategy)

    def test_sentence_strategy_builds_sentence_splitters(self):
        with patch.dict(os.environ, {"CHUNK_STRATEGY": "sentence"}):
            loader = DocumentLoader()
        self.assertEqual("sentence", loader._strategy)
        self.assertIsInstance(loader._splitter_level_3, SentenceSplitter)

    def test_token_strategy_builds_token_splitters_when_tiktoken_present(self):
        try:
            import tiktoken  # noqa: F401
        except ImportError:
            self.skipTest("tiktoken not installed")
        from langchain_text_splitters import TokenTextSplitter

        with patch.dict(os.environ, {"CHUNK_STRATEGY": "token"}):
            loader = DocumentLoader()
        self.assertEqual("token", loader._strategy)
        self.assertIsInstance(loader._splitter_level_3, TokenTextSplitter)

    def test_explicit_constructor_args_beat_env(self):
        with patch.dict(os.environ, {"CHUNK_SIZE": "1000"}):
            loader = DocumentLoader(chunk_size=800)
        self.assertEqual(800, loader._level_3_size)


class SentenceSplitterTests(unittest.TestCase):
    def test_packs_sentences_within_budget_with_sentence_overlap(self):
        splitter = SentenceSplitter(max_chars=60, overlap_sentences=1)
        text = "First sentence here. Second sentence there. Third sentence now."
        chunks = splitter.split_text(text)
        self.assertGreater(len(chunks), 1)
        # Overlap is whole sentences: each next chunk starts with the previous tail.
        self.assertTrue(chunks[1].startswith(chunks[0].split(". ")[-1].split(".")[0][:10]))
        for chunk in chunks:
            self.assertLessEqual(len(chunk), 60 + 40)

    def test_single_oversized_sentence_stays_intact(self):
        splitter = SentenceSplitter(max_chars=20, overlap_sentences=0)
        text = "This single sentence is far longer than the twenty character budget."
        self.assertEqual([text], splitter.split_text(text))


class DocxLayoutTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        import docx as docx_lib

        cls._tmp = tempfile.TemporaryDirectory()
        cls.path = str(Path(cls._tmp.name) / "handbook.docx")
        document = docx_lib.Document()
        document.add_heading("Admission Fees", level=1)
        document.add_paragraph("Fees are reviewed annually by the board.")
        table = document.add_table(rows=3, cols=2)
        for row, values in zip(table.rows, [["Grade", "Fee"], ["1", "100"], ["2", "200"]]):
            for cell, value in zip(row.cells, values):
                cell.text = value
        document.add_heading("Transportation", level=1)
        document.add_paragraph("Bus service is optional for all grades.")
        document.save(cls.path)

    @classmethod
    def tearDownClass(cls):
        cls._tmp.cleanup()

    def test_blocks_carry_headings_text_and_real_tables(self):
        blocks = parse_docx_blocks(self.path)
        types = [block["type"] for block in blocks]
        self.assertEqual(["heading", "text", "table", "heading", "text"], types)
        self.assertEqual(1, blocks[0]["level"])
        self.assertEqual(
            [["Grade", "Fee"], ["1", "100"], ["2", "200"]],
            blocks[2]["rows"],
        )

    def test_loader_end_to_end_word_table_is_row_safe_and_topic_prefixed(self):
        docs = DocumentLoader().load_document(self.path, "handbook.docx")
        table_leaf = next(
            d for d in docs if d["chunk_level"] == 3 and "Grade | Fee" in d["text"]
        )
        self.assertTrue(table_leaf["text"].startswith("Admission Fees\n"))
        self.assertIn("2 | 200", table_leaf["text"])
        self.assertNotIn("Bus service", table_leaf["text"])
        self.assertEqual("Word", table_leaf["file_type"])
        transport_leaf = next(
            d for d in docs if d["chunk_level"] == 3 and "Bus service" in d["text"]
        )
        self.assertTrue(transport_leaf["text"].startswith("Transportation"))

    def test_legacy_doc_extension_never_touches_layout_parser(self):
        class _StubDoc:
            page_content = "legacy word text"
            metadata = {}

        class _StubDocx2txtLoader:
            def __init__(self, file_path):
                pass

            def load(self):
                return [_StubDoc()]

        layout_mock = Mock()
        with patch.object(document_loader_module, "parse_docx_blocks", layout_mock), patch.object(
            document_loader_module, "Docx2txtLoader", _StubDocx2txtLoader
        ):
            docs = DocumentLoader().load_document("legacy.doc", "legacy.doc")
        layout_mock.assert_not_called()
        self.assertTrue(docs)


class XlsxLayoutTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        from openpyxl import Workbook

        cls._tmp = tempfile.TemporaryDirectory()
        cls.path = str(Path(cls._tmp.name) / "fees.xlsx")
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Tuition"
        for row in (["Grade", "Fee"], ["1", "100"], ["2", "200"]):
            sheet.append(row)
        workbook.save(cls.path)

    @classmethod
    def tearDownClass(cls):
        cls._tmp.cleanup()

    def test_sheet_becomes_heading_plus_table_blocks(self):
        blocks = parse_xlsx_blocks(self.path)
        self.assertEqual(["heading", "table"], [block["type"] for block in blocks])
        self.assertEqual("Tuition", blocks[0]["content"])
        self.assertEqual([["Grade", "Fee"], ["1", "100"], ["2", "200"]], blocks[1]["rows"])

    def test_loader_end_to_end_sheet_table_under_sheet_topic(self):
        docs = DocumentLoader().load_document(self.path, "fees.xlsx")
        table_leaf = next(
            d for d in docs if d["chunk_level"] == 3 and "Grade | Fee" in d["text"]
        )
        self.assertTrue(table_leaf["text"].startswith("Tuition\n"))
        self.assertEqual("Excel", table_leaf["file_type"])


class HtmlLayoutTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls._tmp = tempfile.TemporaryDirectory()
        cls.path = str(Path(cls._tmp.name) / "guide.html")
        Path(cls.path).write_text(
            """
            <html><head><title>School Guide</title></head><body><main>
            <h1>Fees</h1>
            <p>Annual tuition depends on grade level.</p>
            <table>
              <tr><th>Grade</th><th>Fee</th></tr>
              <tr><td>1</td><td>100</td></tr>
              <tr><td>2</td><td>200</td></tr>
            </table>
            <h2>Notes</h2>
            <p>Discounts apply for siblings enrolled together.</p>
            </main></body></html>
            """,
            encoding="utf-8",
        )

    @classmethod
    def tearDownClass(cls):
        cls._tmp.cleanup()

    def test_blocks_carry_title_headings_and_real_table(self):
        blocks = parse_html_blocks(self.path)
        headings = [b["content"] for b in blocks if b["type"] == "heading"]
        self.assertEqual(["School Guide", "Fees", "Notes"], headings)
        table = next(b for b in blocks if b["type"] == "table")
        self.assertEqual([["Grade", "Fee"], ["1", "100"], ["2", "200"]], table["rows"])
        # Table cell text must not leak into text blocks.
        for block in blocks:
            if block["type"] == "text":
                self.assertNotIn("100", block["content"])

    def test_loader_end_to_end_html_table_under_section(self):
        docs = DocumentLoader().load_document(self.path, "guide.html")
        table_leaf = next(
            d for d in docs if d["chunk_level"] == 3 and "Grade | Fee" in d["text"]
        )
        self.assertTrue(table_leaf["text"].startswith("Fees\n"))
        self.assertEqual("HTML", table_leaf["file_type"])


class HtmlEdgeCaseTests(unittest.TestCase):
    """Error guessing on HTML structure: layout tables, nested tables."""

    def _parse(self, body):
        tmp = tempfile.TemporaryDirectory()
        self.addCleanup(tmp.cleanup)
        path = str(Path(tmp.name) / "page.html")
        Path(path).write_text(f"<html><body><main>{body}</main></body></html>", encoding="utf-8")
        return parse_html_blocks(path)

    def test_layout_table_with_prose_cells_is_demoted_to_text(self):
        blocks = self._parse(
            "<table><tr><td>This is a long prose sentence living inside a layout table cell.</td></tr>"
            "<tr><td>Another long prose sentence used purely for page layout purposes.</td></tr></table>"
        )
        self.assertTrue(all(block["type"] == "text" for block in blocks))

    def test_nested_table_rows_are_not_double_counted(self):
        blocks = self._parse(
            "<table>"
            "<tr><th>Grade</th><th>Fee</th></tr>"
            "<tr><td>1</td><td>100</td></tr>"
            "<tr><td>2</td><td><table><tr><td>inner1</td><td>x</td></tr>"
            "<tr><td>inner2</td><td>y</td></tr></table></td></tr>"
            "</table>"
        )
        tables = [b for b in blocks if b["type"] == "table"]
        self.assertEqual(1, len(tables))
        # Outer table has exactly its own 3 rows; inner rows must not appear as
        # extra top-level rows.
        self.assertEqual(3, len(tables[0]["rows"]))

    def test_list_items_become_text_blocks(self):
        blocks = self._parse("<ul><li>Bring the vaccination card.</li><li>Bring two photos.</li></ul>")
        texts = [b["content"] for b in blocks if b["type"] == "text"]
        self.assertIn("Bring the vaccination card.", texts)
        self.assertIn("Bring two photos.", texts)


class XlsxEdgeCaseTests(unittest.TestCase):
    """EP: multi-sheet isolation and empty-sheet partition."""

    @classmethod
    def setUpClass(cls):
        from openpyxl import Workbook

        cls._tmp = tempfile.TemporaryDirectory()
        cls.path = str(Path(cls._tmp.name) / "multi.xlsx")
        workbook = Workbook()
        fees = workbook.active
        fees.title = "Fees"
        for row in (["Grade", "Fee"], ["1", "100"]):
            fees.append(row)
        buses = workbook.create_sheet("Buses")
        for row in (["Route", "Cost"], ["North", "50"]):
            buses.append(row)
        workbook.create_sheet("EmptySheet")
        workbook.save(cls.path)

    @classmethod
    def tearDownClass(cls):
        cls._tmp.cleanup()

    def test_empty_sheet_produces_no_blocks(self):
        blocks = parse_xlsx_blocks(self.path)
        self.assertNotIn("EmptySheet", [b["content"] for b in blocks if b["type"] == "heading"])

    def test_same_column_count_sheets_are_never_stitched_together(self):
        docs = DocumentLoader().load_document(self.path, "multi.xlsx")
        fees_leaf = next(d for d in docs if d["chunk_level"] == 3 and "Grade | Fee" in d["text"])
        buses_leaf = next(d for d in docs if d["chunk_level"] == 3 and "Route | Cost" in d["text"])
        self.assertTrue(fees_leaf["text"].startswith("Fees\n"))
        self.assertTrue(buses_leaf["text"].startswith("Buses\n"))
        self.assertNotIn("Route", fees_leaf["text"])
        self.assertNotIn("Grade", buses_leaf["text"])


class DocxEdgeCaseTests(unittest.TestCase):
    """EP: heading style depth, empty paragraphs."""

    @classmethod
    def setUpClass(cls):
        import docx as docx_lib

        cls._tmp = tempfile.TemporaryDirectory()
        cls.path = str(Path(cls._tmp.name) / "styles.docx")
        document = docx_lib.Document()
        document.add_heading("Chapter", level=1)
        document.add_paragraph("")  # empty paragraph must be skipped
        document.add_heading("Deep Section", level=3)
        document.add_paragraph("Content under the deep section heading.")
        document.save(cls.path)

    @classmethod
    def tearDownClass(cls):
        cls._tmp.cleanup()

    def test_heading_levels_and_empty_paragraph_skip(self):
        blocks = parse_docx_blocks(self.path)
        self.assertEqual(["heading", "heading", "text"], [b["type"] for b in blocks])
        self.assertEqual(1, blocks[0]["level"])
        self.assertEqual(3, blocks[1]["level"])


if __name__ == "__main__":
    unittest.main()
